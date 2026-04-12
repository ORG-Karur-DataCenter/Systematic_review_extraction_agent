"""
Generate OSA in Spine Surgery Data Extraction Template (Wide Format)
Each outcome = its own column. Total ~33 columns across main sheets.
"""

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document

# ────────────────────────────────────────────
# SHEET 1: STUDY CHARACTERISTICS (15 columns)
# ────────────────────────────────────────────
CHARACTERISTICS = [
    ("Study ID", "First Author, Year (e.g., Patel 2024)"),
    ("Study Design", "Retrospective cohort / Prospective cohort / Cross-sectional / Case-control / RCT / Database study"),
    ("Country", "Country of study"),
    ("Data Source", "Single-center / NIS / ACS-NSQIP / Other"),
    ("Study Period", "Start - End year"),
    ("Spine Region & Procedure", "e.g., Lumbar fusion (TLIF), Cervical ACDF, Decompression"),
    ("Total N", "Total sample size"),
    ("OSA n", "Patients with OSA"),
    ("Control n", "Patients without OSA (N/A if single-arm)"),
    ("OSA Diagnostic Method", "PSG / ICD codes / STOP-BANG / Berlin / HSAT / Self-report"),
    ("OSA Severity (Mild/Mod/Severe n)", "AHI-based breakdown or NR"),
    ("Age, Sex, BMI", "Report for OSA vs Control: e.g., 58.2+/-7.1 vs 54.3+/-8.0; 62% vs 55% male; BMI 32.1 vs 27.4"),
    ("Key Comorbidities", "HTN, DM, COPD, CVD, CCI - rates for OSA vs Control"),
    ("CPAP Use / Compliance", "Pre-op CPAP %, compliance, post-op use; or NR"),
    ("Follow-up Duration", "Mean/Median follow-up period"),
]

# ────────────────────────────────────────────
# SHEET 2: OUTCOMES (18 columns, one per outcome)
# Each cell format: "OSA: n/N (%), Control: n/N (%), OR/aOR X.X (CI), p=X.XX"
# For single-arm: just "n/N (%)"
# For continuous: "OSA: Mean+/-SD, Control: Mean+/-SD, MD X.X (CI), p=X.XX"
# ────────────────────────────────────────────
OUTCOMES = [
    ("Study ID", "Must match Sheet 1"),
    ("Pneumonia", "n/N (%) per group, effect estimate, p-value"),
    ("Respiratory Failure", "n/N (%) per group, effect estimate, p-value"),
    ("Reintubation", "Unplanned reintubation: n/N (%) per group, effect estimate, p-value"),
    ("Desaturation Events", "n/N (%) or events/patient per group, effect estimate"),
    ("Cardiac Complications", "Arrhythmia, MI, cardiac arrest: n/N (%) per group, effect estimate"),
    ("DVT / PE", "Deep vein thrombosis or pulmonary embolism: n/N (%) per group, effect estimate"),
    ("SSI", "Surgical site infection: n/N (%) per group, effect estimate"),
    ("Wound Complications", "Dehiscence, hematoma: n/N (%) per group, effect estimate"),
    ("Neurological Deficit", "New neurological deficit: n/N (%) per group, effect estimate"),
    ("AKI", "Acute kidney injury: n/N (%) per group, effect estimate"),
    ("Sepsis / UTI", "Sepsis or urinary tract infection: n/N (%) per group, effect estimate"),
    ("Blood Transfusion", "n/N (%) per group, effect estimate"),
    ("LOS (days)", "Length of stay: Mean+/-SD or Median(IQR) per group, mean difference, p-value"),
    ("ICU Admission", "n/N (%) or ICU LOS per group, effect estimate"),
    ("Readmission (30d/90d)", "n/N (%) per group, effect estimate"),
    ("Mortality", "In-hospital or 30-day: n/N (%) per group, effect estimate, p-value"),
    ("Other Outcomes", "PROs (VAS, ODI, NDI), sleep outcomes (AHI change, ESS), costs, or any other reported outcome"),
]

# ────────────────────────────────────────────
# SHEET 3: QUALITY ASSESSMENT (7 columns)
# ────────────────────────────────────────────
QUALITY = [
    ("Study ID", "Must match Sheet 1"),
    ("RoB Tool", "NOS / MINORS / RoB 2.0 / ROBINS-I"),
    ("Selection Score", "NOS: 0-4 stars; or Low/Moderate/High"),
    ("Comparability Score", "NOS: 0-2 stars; or Low/Moderate/High"),
    ("Outcome Score", "NOS: 0-3 stars; or Low/Moderate/High"),
    ("Overall Score / Rating", "NOS total (0-9) or Overall RoB"),
    ("Key Limitations", "Major study limitations"),
]


def create_xlsx():
    wb = openpyxl.Workbook()

    hdr_font = Font(name="Calibri", bold=True, size=11, color="FFFFFF")
    hdr_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    desc_font = Font(name="Calibri", italic=True, size=9, color="808080")
    thin = Border(
        left=Side("thin", "B4C6E7"), right=Side("thin", "B4C6E7"),
        top=Side("thin", "B4C6E7"), bottom=Side("thin", "B4C6E7"),
    )

    def write_sheet(ws, title, fields):
        ws.title = title
        for col_idx, (name, desc) in enumerate(fields, 1):
            c = ws.cell(row=1, column=col_idx, value=name)
            c.font = hdr_font
            c.fill = hdr_fill
            c.alignment = Alignment(horizontal="center", wrap_text=True)
            c.border = thin
            d = ws.cell(row=2, column=col_idx, value=desc)
            d.font = desc_font
            d.alignment = Alignment(wrap_text=True)
            d.border = thin
        for col_idx in range(1, len(fields) + 1):
            ws.column_dimensions[get_column_letter(col_idx)].width = 24
        ws.freeze_panes = "A3"
        for row in range(3, 25):
            for col_idx in range(1, len(fields) + 1):
                ws.cell(row=row, column=col_idx).border = thin

    ws1 = wb.active
    write_sheet(ws1, "Characteristics", CHARACTERISTICS)
    write_sheet(wb.create_sheet(), "Outcomes", OUTCOMES)
    write_sheet(wb.create_sheet(), "Quality Assessment", QUALITY)

    out = "OSA_Spine_Surgery_Data_Extraction_Template.xlsx"
    wb.save(out)
    print(f"[OK] Excel: {out}")
    print(f"     Characteristics: {len(CHARACTERISTICS)} cols | Outcomes: {len(OUTCOMES)} cols | Quality: {len(QUALITY)} cols")
    print(f"     TOTAL: {len(CHARACTERISTICS) + len(OUTCOMES) + len(QUALITY)} columns")


def create_docx():
    doc = Document()
    doc.add_heading("OSA in Spine Surgery - Data Extraction Template", level=0)

    doc.add_heading("Study Characteristics", level=1)
    for name, desc in CHARACTERISTICS:
        doc.add_paragraph(f"{name}: {desc}")

    doc.add_heading("Outcomes", level=1)
    for name, desc in OUTCOMES:
        doc.add_paragraph(f"{name}: {desc}")

    doc.add_heading("Quality Assessment", level=1)
    for name, desc in QUALITY:
        doc.add_paragraph(f"{name}: {desc}")

    out = "OSA_Spine_Surgery_Data_Extraction_Template.docx"
    doc.save(out)
    print(f"[OK] DOCX: {out}")


if __name__ == "__main__":
    create_xlsx()
    create_docx()

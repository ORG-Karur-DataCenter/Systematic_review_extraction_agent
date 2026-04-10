"""
Template Parser Module

This module provides a unified interface for parsing data extraction templates
in both Word (.docx) and Excel (.xlsx) formats. It automatically detects the
template format and extracts field definitions for use in the extraction agent.

Usage:
    from template_parser import parse_template
    
    fields = parse_template('template.docx')
    # or
    fields = parse_template('template.xlsx')
"""

import os
import re
from typing import List, Dict, Optional
import docx
import pandas as pd


class TemplateField:
    """Represents a single field in the extraction template."""
    
    def __init__(self, name: str, description: str = "", section: str = ""):
        self.name = name.strip()
        self.description = description.strip()
        self.section = section.strip()
    
    def to_dict(self) -> Dict[str, str]:
        """Convert field to dictionary representation."""
        return {
            'name': self.name,
            'description': self.description,
            'section': self.section
        }
    
    def __repr__(self):
        return f"TemplateField(name='{self.name}', section='{self.section}')"


class TemplateParser:
    """Base class for template parsers."""
    
    def parse(self, filepath: str) -> List[TemplateField]:
        """Parse template and return list of fields."""
        raise NotImplementedError


class WordTemplateParser(TemplateParser):
    """Parser for Word (.docx) templates."""
    
    # Section headers typically don't end with colons
    SECTION_PATTERNS = [
        r'^[A-Z][A-Za-z\s]+(?:Details|Characteristics|Outcomes|Identification|Information)$',
        r'^[A-Z][A-Za-z\s]+\([A-Za-z\s:±]+\)$',  # e.g., "Baseline Characteristics (Continuous: Mean ± SD)"
    ]
    
    # Field patterns typically end with colons
    FIELD_PATTERN = r'^(.+?):\s*(.*)$'
    
    def parse(self, filepath: str) -> List[TemplateField]:
        """Parse Word template and extract fields."""
        doc = docx.Document(filepath)
        fields = []
        current_section = ""
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this is a section header
            if self._is_section_header(text):
                current_section = text
                continue
            
            # Check if this is a field definition
            field_match = re.match(self.FIELD_PATTERN, text)
            if field_match:
                field_name = field_match.group(1).strip()
                field_desc = field_match.group(2).strip()
                
                # Skip if it looks like a title
                if self._is_likely_title(field_name):
                    continue
                
                fields.append(TemplateField(
                    name=field_name,
                    description=field_desc,
                    section=current_section
                ))
        
        # Also check tables in Word document
        for table in doc.tables:
            table_fields = self._parse_table(table, current_section)
            fields.extend(table_fields)
        
        return fields
    
    def _is_section_header(self, text: str) -> bool:
        """Determine if text is a section header."""
        # Section headers typically don't have colons at the end
        if text.endswith(':'):
            return False
        
        # Check against known patterns
        for pattern in self.SECTION_PATTERNS:
            if re.match(pattern, text):
                return True
        
        # Check if it's all caps or title case without colon
        if text.isupper() or (text.istitle() and len(text.split()) <= 6):
            return True
        
        return False
    
    def _is_likely_title(self, text: str) -> bool:
        """Check if text is likely a document title rather than a field."""
        title_keywords = ['template', 'data extraction', 'meta-analysis']
        return any(keyword in text.lower() for keyword in title_keywords)
    
    def _parse_table(self, table, current_section: str) -> List[TemplateField]:
        """Parse fields from a Word table."""
        fields = []
        
        # Assume first row might be headers
        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:
                # Check if this looks like a header row
                first_cell = row.cells[0].text.strip().lower()
                if 'field' in first_cell or 'column' in first_cell:
                    continue
            
            # Extract field name and description from columns
            if len(row.cells) >= 2:
                field_name = row.cells[0].text.strip()
                field_desc = row.cells[1].text.strip()
                
                if field_name and not self._is_likely_title(field_name):
                    fields.append(TemplateField(
                        name=field_name,
                        description=field_desc,
                        section=current_section
                    ))
        
        return fields


class ExcelTemplateParser(TemplateParser):
    """Parser for Excel (.xlsx) templates."""
    
    def parse(self, filepath: str) -> List[TemplateField]:
        """Parse Excel template and extract fields from column headers."""
        xl = pd.ExcelFile(filepath)
        fields = []
        
        # Parse all sheets
        for sheet_name in xl.sheet_names:
            df = pd.read_excel(xl, sheet_name=sheet_name)
            
            # Column names are the field names
            for col_name in df.columns:
                # Skip unnamed columns
                if isinstance(col_name, str) and not col_name.startswith('Unnamed'):
                    fields.append(TemplateField(
                        name=col_name,
                        description="",  # Excel templates typically don't have descriptions
                        section=sheet_name
                    ))
        
        return fields


def detect_template_format(filepath: str) -> Optional[str]:
    """
    Detect template format based on file extension.
    
    Args:
        filepath: Path to template file
    
    Returns:
        'word' for .docx files, 'excel' for .xlsx files, None if unknown
    """
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == '.docx':
        return 'word'
    elif ext in ['.xlsx', '.xls']:
        return 'excel'
    else:
        return None


def parse_template(filepath: str) -> List[TemplateField]:
    """
    Parse a template file and extract field definitions.
    
    Automatically detects whether the template is in Word or Excel format
    and uses the appropriate parser.
    
    Args:
        filepath: Path to template file (.docx or .xlsx)
    
    Returns:
        List of TemplateField objects
    
    Raises:
        FileNotFoundError: If template file doesn't exist
        ValueError: If template format is not supported
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"Template file not found: {filepath}")
    
    format_type = detect_template_format(filepath)
    
    if format_type == 'word':
        parser = WordTemplateParser()
    elif format_type == 'excel':
        parser = ExcelTemplateParser()
    else:
        raise ValueError(f"Unsupported template format: {filepath}")
    
    fields = parser.parse(filepath)
    
    if not fields:
        raise ValueError(f"No fields found in template: {filepath}")
    
    return fields


def fields_to_dict_list(fields: List[TemplateField]) -> List[Dict[str, str]]:
    """Convert list of TemplateField objects to list of dictionaries."""
    return [field.to_dict() for field in fields]


def get_field_names(fields: List[TemplateField]) -> List[str]:
    """Extract just the field names from a list of TemplateField objects."""
    return [field.name for field in fields]


if __name__ == "__main__":
    # Example usage and testing
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python template_parser.py <template_file>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    
    try:
        print(f"Parsing template: {template_path}")
        print(f"Format: {detect_template_format(template_path)}")
        print()
        
        fields = parse_template(template_path)
        
        print(f"Found {len(fields)} fields:")
        print("=" * 80)
        
        current_section = ""
        for field in fields:
            if field.section != current_section:
                current_section = field.section
                print(f"\n[{current_section}]")
            
            if field.description:
                print(f"  - {field.name}: {field.description}")
            else:
                print(f"  - {field.name}")
        
        print("\n" + "=" * 80)
        print(f"Total fields: {len(fields)}")
        
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)
